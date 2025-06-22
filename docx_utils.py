import json
from docx import Document

def generate_pretty_docx(json_data_str, output_path="generated_report.docx"):
    try:
        data = json.loads(json_data_str)
    except json.JSONDecodeError:
        doc = Document()
        doc.add_heading("Relatório Clínico", level=1)
        doc.add_paragraph("Erro: JSON inválido.")
        doc.save(output_path)
        return

    doc = Document()
    doc.add_heading("Relatório Clínico", level=1)

    def add_section(title, content):
        doc.add_heading(title, level=2)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style="List Bullet")
        elif isinstance(content, dict):
            for k, v in content.items():
                doc.add_paragraph(f"{k.capitalize()}: {v}")
        elif content:
            doc.add_paragraph(str(content))

    doc.add_paragraph(f"Data do Exame: {data.get('data_exame', '')}")
    doc.add_paragraph(f"Clínica: {data.get('clinica', '')}")

    animal = data.get("animal", {})
    doc.add_heading("Animal", level=2)
    for k, v in animal.items():
        if v:  # only include non-empty fields
            doc.add_paragraph(f"{k.capitalize()}: {v}")

    add_section("Motivo da Consulta", data.get("motivo_consulta", ""))
    add_section("Anamnese", data.get("anamnese", ""))
    add_section("Exame Estático", data.get("exame_estatico", {}).get("observacoes", []))
    add_section("Exame Dinâmico", data.get("exame_dinamico", {}).get("observacoes", []))

    exames = data.get("exames_complementares", {})
    if exames:
        doc.add_heading("Exames Complementares", level=2)
        for tipo, sub in exames.items():
            for nome, detalhe in sub.items():
                doc.add_paragraph(f"{tipo.upper()} - {nome}", style="List Bullet")
                doc.add_paragraph(f"Descrição: {detalhe.get('descricao', '')}")
                for proj in detalhe.get("projecoes", []):
                    doc.add_paragraph(f"Projeção: {proj}", style="List Bullet")
                for achado in detalhe.get("achados", []):
                    doc.add_paragraph(f"Achado: {achado}", style="List Bullet")

    tratamento = data.get("tratamento", [])
    if tratamento:
        add_section("Tratamento", [f"{t['tipo']} - {t['medicamento']} ({t['dose']})" for t in tratamento])

    add_section("Recomendações", data.get("recomendacoes", {}).get("observacoes_adicionais", []))
    add_section("Indicações de Ferragem", data.get("indicacoes_ferragem", {}))

    doc.save(output_path)
